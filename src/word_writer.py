import logging
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src import PageResult, TableData
from src.layout import (
    indent_for_line,
    page_text_bounds,
    sort_lines,
    space_before_for_line,
    estimate_font_size,
    detect_alignment,
    estimate_is_bold,
    ALIGN_LEFT,
    ALIGN_CENTER,
    ALIGN_RIGHT,
)

_SYMBOL_MAP = {
    r"\times": "×",
    r"\cdot": "·",
    r"\geq": "≥",
    r"\leq": "≤",
    r"\sim": "~",
    r"\quad": " ",
    r"\pm": "±",
}

log = logging.getLogger(__name__)

_BRACED_RE = re.compile(r"_\{|\^\{|\\frac\{")
_LARGE_BRACES = {"{", "}"}


def _color(score: float):
    if score < 0.5:
        return RGBColor(0xCC, 0x00, 0x00)
    if score < 0.65:
        return RGBColor(0xCC, 0x66, 0x00)
    if score < 0.75:
        return RGBColor(0x99, 0x66, 0x00)
    return None


def _set_run_font(run, font_cfg: dict, size_pt: float | None = None):
    latin = font_cfg.get("latin", "Times New Roman")
    ea = font_cfg.get("east_asian", "宋体")
    size = size_pt if size_pt is not None else font_cfg.get("size", 10)
    run.font.name = latin
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    rFonts.set(qn("w:eastAsia"), ea)


def _replace_symbols(text: str) -> str:
    for src, dst in _SYMBOL_MAP.items():
        text = text.replace(src, dst)
    text = re.sub(r"\\text\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+|[{}]", "", text)
    return text.strip()


def _parse_braced(text: str, start: int) -> tuple[str, int]:
    if start >= len(text) or text[start] != "{":
        return "", start
    depth = 1
    i = start + 1
    while i < len(text) and depth > 0:
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
        i += 1
    return text[start + 1 : i - 1], i


def _parse_fraction(text: str, start: int) -> tuple[str, str, int]:
    if text[start : start + 6] != r"\frac{":
        return "", "", start
    num, i = _parse_braced(text, start + 5)
    den, i = _parse_braced(text, i)
    return num, den, i


def _make_math_para(omml_elem) -> OxmlElement:
    mp = OxmlElement("m:oMathPara")
    mp.append(omml_elem)
    return mp


def _make_run_elem(text: str, font_cfg: dict) -> OxmlElement:
    r = OxmlElement("m:r")
    rPr = OxmlElement("m:rPr")
    latin = font_cfg.get("latin", "Times New Roman")
    ea = font_cfg.get("east_asian", "宋体")
    size = font_cfg.get("size", 10)
    rFonts = OxmlElement("m:rFonts")
    rFonts.set(qn("m:ascii"), latin)
    rFonts.set(qn("m:hAnsi"), latin)
    rFonts.set(qn("m:eastAsia"), ea)
    rPr.append(rFonts)
    sz = OxmlElement("m:sz")
    sz.set(qn("m:val"), str(size * 2))
    rPr.append(sz)
    szCs = OxmlElement("m:szCs")
    szCs.set(qn("m:val"), str(size * 2))
    rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement("m:t")
    t.text = text
    r.append(t)
    return r


def _build_fraction(num_text: str, den_text: str, font_cfg: dict) -> OxmlElement:
    oMath = OxmlElement("m:oMath")
    f = OxmlElement("m:f")
    fPr = OxmlElement("m:fPr")
    typ = OxmlElement("m:type")
    typ.set(qn("m:val"), "bar")
    fPr.append(typ)
    f.append(fPr)

    num = OxmlElement("m:num")
    num.append(_make_run_elem(_replace_symbols(num_text), font_cfg))
    f.append(num)

    den = OxmlElement("m:den")
    den.append(_make_run_elem(_replace_symbols(den_text), font_cfg))
    f.append(den)

    oMath.append(f)
    return oMath


def _build_subscript(base: str, sub: str, font_cfg: dict) -> OxmlElement:
    oMath = OxmlElement("m:oMath")
    sSub = OxmlElement("m:sSub")
    sSubPr = OxmlElement("m:sSubPr")
    sSubPr.append(OxmlElement("m:ctrlPr"))
    sSub.append(sSubPr)

    e = OxmlElement("m:e")
    e.append(_make_run_elem(_replace_symbols(base), font_cfg))
    sSub.append(e)

    sub_elem = OxmlElement("m:sub")
    sub_elem.append(_make_run_elem(_replace_symbols(sub), font_cfg))
    sSub.append(sub_elem)

    oMath.append(sSub)
    return oMath


def _build_superscript(base: str, sup: str, font_cfg: dict) -> OxmlElement:
    oMath = OxmlElement("m:oMath")
    sSup = OxmlElement("m:sSup")
    sSupPr = OxmlElement("m:sSupPr")
    sSupPr.append(OxmlElement("m:ctrlPr"))
    sSup.append(sSupPr)

    e = OxmlElement("m:e")
    e.append(_make_run_elem(_replace_symbols(base), font_cfg))
    sSup.append(e)

    sup_elem = OxmlElement("m:sup")
    sup_elem.append(_make_run_elem(_replace_symbols(sup), font_cfg))
    sSup.append(sup_elem)

    oMath.append(sSup)
    return oMath


def _add_text_run(para, text: str, font_cfg: dict, size_pt: float | None = None):
    text = _replace_symbols(text)
    if not text:
        return
    run = para.add_run(text)
    _set_run_font(run, font_cfg, size_pt)


def _add_math_content(para, text: str, font_cfg: dict, size_pt: float | None = None):
    text = re.sub(r"\\\(|\\\)", "", text)
    text = re.sub(r"\\\[|\\\]", "", text)

    while text:
        m_frac = re.search(r"\\frac\{", text)
        m_sub = re.search(r"[_^]", text)

        if m_frac is not None and (m_sub is None or m_frac.start() < m_sub.start()):
            before = text[: m_frac.start()]
            if before:
                _add_text_run(para, before, font_cfg, size_pt)
            num, den, end = _parse_fraction(text, m_frac.start())
            para._element.append(_build_fraction(num, den, font_cfg))
            text = text[end:]
            continue

        if m_sub is not None:
            base_end = m_sub.start()
            base_start = base_end - 1
            while base_start >= 0 and (text[base_start].isalnum() or text[base_start] in "_."):
                base_start -= 1
            base_start += 1
            base = text[base_start:base_end]

            if base_start > 0:
                _add_text_run(para, text[:base_start], font_cfg, size_pt)

            op = text[m_sub.start()]
            if m_sub.end() < len(text) and text[m_sub.end()] == "{":
                content, end = _parse_braced(text, m_sub.end())
            else:
                content = text[m_sub.end()] if m_sub.end() < len(text) else ""
                end = m_sub.end() + 1

            if op == "_":
                para._element.append(_build_subscript(base, content, font_cfg))
            else:
                para._element.append(_build_superscript(base, content, font_cfg))
            text = text[end:]
            continue

        _add_text_run(para, text, font_cfg, size_pt)
        break


def _add_large_brace(para, text: str, bbox: tuple[int, int, int, int], page_h: int, font_cfg: dict, size_pt: float | None = None):
    run = para.add_run(text)
    _set_run_font(run, font_cfg, size_pt)
    bbox_h = max(bbox[3] - bbox[1], 1)
    ratio = bbox_h / max(page_h, 1)
    size = max(22, min(int(ratio * 180), 72))
    run.font.size = Pt(size)


def _is_large_brace_line(line) -> bool:
    return line.text.strip() in _LARGE_BRACES


def _line_center_y(line) -> float:
    return (line.bbox[1] + line.bbox[3]) / 2


def _build_brace_blocks(lines, page_w: int):
    braces = [ln for ln in lines if _is_large_brace_line(ln)]
    text_lines = [ln for ln in lines if not _is_large_brace_line(ln)]
    blocks = {}
    claimed = set()

    for brace in braces:
        brace_x1, brace_y1, brace_x2, brace_y2 = brace.bbox
        y_pad = max(brace_y2 - brace_y1, 1) * 0.08
        x_limit = min(page_w, brace_x2 + page_w * 0.55)
        candidates = []
        for line in text_lines:
            if id(line) in claimed:
                continue
            cy = _line_center_y(line)
            if cy < brace_y1 - y_pad or cy > brace_y2 + y_pad:
                continue
            if line.bbox[0] < brace_x2 - page_w * 0.02:
                continue
            if line.bbox[0] > x_limit:
                continue
            candidates.append(line)

        if not candidates:
            continue

        candidates.sort(key=lambda ln: (ln.bbox[1], ln.bbox[0]))
        anchor_x = min(ln.bbox[0] for ln in candidates)
        blocks[id(brace)] = {
            "brace": brace,
            "lines": candidates,
            "first": candidates[0],
            "anchor_x": anchor_x,
        }
        for line in candidates:
            claimed.add(id(line))

    line_to_block = {}
    for block in blocks.values():
        for line in block["lines"]:
            line_to_block[id(line)] = block
    return blocks, line_to_block


def _with_anchor_bbox(line, anchor_x: int):
    x1, y1, x2, y2 = line.bbox
    return type(line)(
        text=line.text,
        score=line.score,
        bbox=(anchor_x, y1, max(x2, anchor_x + 1), y2),
        polygon=line.polygon,
    )


def _add_line_content(para, line, page_h: int, font_cfg: dict, size_pt: float | None = None, bold: bool = False):
    normalized_text = line.text.strip()
    has_latex = "\\" in line.text or "_{" in line.text or "^{" in line.text
    if normalized_text in _LARGE_BRACES:
        _add_large_brace(para, normalized_text, line.bbox, page_h, font_cfg, size_pt)
    elif has_latex:
        _add_math_content(para, line.text, font_cfg, size_pt)
    else:
        run = para.add_run(line.text)
        _set_run_font(run, font_cfg, size_pt)
        if bold:
            run.font.bold = True


def _collect_table_line_ids(tables: list[TableData]) -> set[int]:
    """收集属于表格的 OCR line IDs，以便从正文段落中排除"""
    ids: set[int] = set()
    for table in tables:
        for row in table.cells:
            for cell in row:
                if cell.bbox and any(v != 0 for v in cell.bbox):
                    cx1, cy1, cx2, cy2 = cell.bbox
                    # 标记与表格区域重叠的 OCR lines
                    # 这是由 write_to_word 调用者在外部处理的，
                    # 此处仅作为辅助函数
                    pass
    return ids


def _write_table_to_word(doc, table_data: TableData, font_cfg: dict, page_w: int, page_h: int):
    """将 TableData 写入 Word 表格"""
    if not table_data.cells:
        return

    num_rows = table_data.num_rows
    num_cols = table_data.num_cols

    word_table = doc.add_table(rows=num_rows, cols=num_cols)
    word_table.style = "Table Grid"
    word_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    latin = font_cfg.get("latin", "Times New Roman")
    ea = font_cfg.get("east_asian", "宋体")
    fsize = font_cfg.get("size", 10)

    for ri, row in enumerate(table_data.cells):
        for ci, cell_data in enumerate(row):
            cell = word_table.cell(ri, ci)
            text = cell_data.text.strip()

            # 清空默认段落
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text if text else "")
            run.font.name = latin
            run.font.size = Pt(fsize)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), latin)
            rFonts.set(qn("w:hAnsi"), latin)
            rFonts.set(qn("w:eastAsia"), ea)

            # 自动调整字体大小（适应单元格）
            if text and len(text) > 20:
                run.font.size = Pt(max(fsize - 2, 7))

            # 设置段间距
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    # 调整列宽（等宽分配）
    if num_cols > 0 and page_w > 0:
        # 估算可用宽度为页面宽度的 90%
        usable_w_mm = page_w * 25.4 / 300 * 0.85
        col_w_mm = usable_w_mm / num_cols
        for ci in range(num_cols):
            for row in word_table.rows:
                cell = row.cells[ci]
                cell.width = Cm(col_w_mm / 10)

    log.info("    Word 表格: %d×%d 写入完成", num_rows, num_cols)


def write_to_word(
    pages: list[PageResult],
    output_path: str | Path,
    font_cfg: dict | None = None,
    pw_mm: float = 148,
    ph_mm: float = 210,
) -> Path:
    if font_cfg is None:
        font_cfg = {}

    doc = Document()
    sec = doc.sections[0]
    margin_cm = 1.5

    # 使用第一页的实际 PDF 尺寸（如果有）来设置 Word 页面
    if pages and pages[0].width > 0 and pages[0].height > 0:
        dpi = 300  # PDF 导出 DPI，与 main.py --dpi 默认值一致
        pw_mm = pages[0].width * 25.4 / dpi
        ph_mm = pages[0].height * 25.4 / dpi
        # 限制最小/最大尺寸
        pw_mm = max(100.0, min(pw_mm, 420.0))
        ph_mm = max(140.0, min(ph_mm, 600.0))

    sec.page_width = Cm(pw_mm / 10)
    sec.page_height = Cm(ph_mm / 10)
    sec.top_margin = Cm(margin_cm)
    sec.bottom_margin = Cm(margin_cm)
    sec.left_margin = Cm(margin_cm)
    sec.right_margin = Cm(margin_cm)

    usable_w_cm = pw_mm / 10 - margin_cm * 2
    latin = font_cfg.get("latin", "Times New Roman")
    ea = font_cfg.get("east_asian", "宋体")
    fsize = font_cfg.get("size", 10)

    s = doc.styles["Normal"]
    s.font.name = latin
    s.font.size = Pt(fsize)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    s.paragraph_format.line_spacing = Pt(14)
    rPr = s._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    rFonts.set(qn("w:eastAsia"), ea)

    total_lines = 0
    auto_size = font_cfg.get("auto_size", True)
    auto_align = font_cfg.get("auto_align", True)
    auto_bold = font_cfg.get("auto_bold", True)
    enable_table = font_cfg.get("table", True)

    for pi, page in enumerate(pages):
        if pi > 0:
            doc.add_page_break()
        if page.skip_reason:
            p = doc.add_paragraph(f"[{page.skip_reason}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if not page.lines:
            continue

        pw = page.width
        ph = page.height
        lines = sort_lines(page.lines, pw)
        bounds = page_text_bounds(
            [ln for ln in lines if not _is_large_brace_line(ln)],
            pw,
        )

        # ── 识别属于表格的 OCR lines ──
        table_line_ids: set[int] = set()
        if enable_table and page.tables:
            for table_data in page.tables:
                _write_table_to_word(doc, table_data, font_cfg, pw, ph)
                # 将表格区域内的 OCR line 标记为不需要单独输出
                for ri, row in enumerate(table_data.cells):
                    for ci, cell_data in enumerate(row):
                        cx1, cy1, cx2, cy2 = cell_data.bbox
                        if cx2 - cx1 <= 0 and cy2 - cy1 <= 0:
                            continue
                        for ln in lines:
                            lx1, ly1, lx2, ly2 = ln.bbox
                            # y 方向重叠 > 50% 且 x 方向有重叠 → 认为属于表格
                            oy = max(0.0, min(ly2, cy2) - max(ly1, cy1))
                            oh = max(ly2 - ly1, 1)
                            if oy > 0 and oy / oh > 0.4:
                                ox = max(0.0, min(lx2, cx2) - max(lx1, cx1))
                                if ox > 0:
                                    table_line_ids.add(id(ln))

        brace_blocks, line_to_brace_block = _build_brace_blocks(lines, pw)
        skip_brace_ids = set(brace_blocks.keys())

        prev_line = None
        for ln in lines:
            if id(ln) in skip_brace_ids:
                continue
            if id(ln) in table_line_ids:
                continue

            block = line_to_brace_block.get(id(ln))
            layout_line = ln
            if block is not None:
                if ln is block["first"]:
                    layout_line = block["brace"]
                else:
                    layout_line = _with_anchor_bbox(ln, block["anchor_x"])

            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # ── 字体大小（从 bbox 高度估算） ──
            line_font_size = estimate_font_size(ln, ph, fsize) if auto_size else fsize

            # ── 行间距随字体大小调整 ──
            line_spacing_pt = max(line_font_size * 1.4, 14.0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            p.paragraph_format.line_spacing = Pt(line_spacing_pt)

            # ── 对齐方式 ──
            if auto_align:
                alignment = detect_alignment(ln, pw, bounds, lines, ph)
            else:
                alignment = ALIGN_LEFT
            _ALIGN_MAP = {
                ALIGN_LEFT: WD_ALIGN_PARAGRAPH.LEFT,
                ALIGN_CENTER: WD_ALIGN_PARAGRAPH.CENTER,
                ALIGN_RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
            }
            p.alignment = _ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

            # ── 缩进（仅左对齐时使用） ──
            if alignment == ALIGN_LEFT:
                space_before = space_before_for_line(prev_line, layout_line)
                if space_before:
                    p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.left_indent = Cm(
                    indent_for_line(layout_line, bounds, usable_w_cm)
                )
            else:
                # 居中/右对齐：用段前间距替代缩进
                space_before = space_before_for_line(prev_line, layout_line)
                if space_before:
                    p.paragraph_format.space_before = Pt(space_before)

            # ── 加粗估算 ──
            is_bold = False
            if auto_bold:
                is_bold = estimate_is_bold(ln, pw)
            # 大号字体自动加粗（标题）
            if line_font_size >= fsize * 1.35:
                is_bold = True

            # ── 写入内容 ──
            if block is not None and ln is block["first"]:
                p.paragraph_format.line_spacing = Pt(max(line_spacing_pt, 16))
                _add_large_brace(p, block["brace"].text.strip(), block["brace"].bbox, ph, font_cfg, line_font_size)
                p.add_run("  ")
                _add_line_content(p, ln, ph, font_cfg, line_font_size, is_bold)
            else:
                _add_line_content(p, ln, ph, font_cfg, line_font_size, is_bold)

            # ── 置信度颜色 ──
            c = _color(ln.score)
            if c:
                for run in p.runs:
                    run.font.color.rgb = c

            prev_line = _with_anchor_bbox(ln, block["anchor_x"]) if block is not None else layout_line
            total_lines += 1

    out_path = Path(output_path)
    doc.save(str(out_path))

    # 完整性校验
    total_input_lines = sum(len(p.lines) for p in pages if not p.skip_reason)
    total_chars = sum(len(ln.text) for p in pages for ln in p.lines)
    log.info(
        "Word 生成完成: %d 页, %d 行文本, %d 字符, 页面 %.1f×%.1f mm → %s",
        len(pages), total_input_lines, total_chars, pw_mm, ph_mm, out_path,
    )
    return out_path
