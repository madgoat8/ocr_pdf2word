#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
验证表格检测功能：
1. table_detector 导入
2. 空间分析表格检测（模拟 OCR lines）
3. 空表格过滤
"""
import logging
import sys

logging.basicConfig(level=logging.INFO, format="%(message)s")

from src import OCRLine, CellData, TableData, PageResult
from src.table_detector import (
    detect_tables_spatial,
    detect_tables_from_pdf,
    detect_tables,
    _group_into_rows,
    _detect_column_boundaries,
    _build_table_from_grid,
)

print("=" * 70)
print("测试 1: 数据结构导入")
print("=" * 70)
print(f"  CellData:    OK")
print(f"  TableData:   OK")
print(f"  PageResult  (has tables): {'tables' in dir(PageResult)}")
print("  [PASS]")

print("\n" + "=" * 70)
print("测试 2: 空间分析 - 行分组")
print("=" * 70)

PAGE_W, PAGE_H = 2000, 3000

# 模拟表格数据：3 行，每行 3 个单元格
lines = [
    OCRLine(text="序号", score=0.9, bbox=(100, 100, 300, 140), polygon=[]),
    OCRLine(text="名称", score=0.9, bbox=(500, 100, 1000, 140), polygon=[]),
    OCRLine(text="备注", score=0.9, bbox=(1200, 100, 1500, 140), polygon=[]),
    OCRLine(text="1", score=0.9, bbox=(100, 200, 300, 240), polygon=[]),
    OCRLine(text="设备A", score=0.9, bbox=(500, 200, 1000, 240), polygon=[]),
    OCRLine(text="正常", score=0.9, bbox=(1200, 200, 1500, 240), polygon=[]),
    OCRLine(text="2", score=0.9, bbox=(100, 300, 300, 340), polygon=[]),
    OCRLine(text="设备B", score=0.9, bbox=(500, 300, 1000, 340), polygon=[]),
    OCRLine(text="异常", score=0.9, bbox=(1200, 300, 1500, 340), polygon=[]),
]

rows = _group_into_rows(lines)
print(f"  行数: {len(rows)} (预期: 3)")
assert len(rows) == 3, f"预期 3 行, 实际 {len(rows)}"
for ri, row in enumerate(rows):
    print(f"    行 {ri+1}: {len(row)} 个单元格 - {[ln.text for ln in row]}")
    assert len(row) == 3, f"行 {ri+1} 应有 3 个单元格, 实际 {len(row)}"
print("  [PASS]")

print("\n" + "=" * 70)
print("测试 3: 空间分析 - 列边界检测")
print("=" * 70)
boundaries = _detect_column_boundaries(rows, PAGE_W, PAGE_H)
print(f"  列边界: {[f'{b:.0f}' for b in boundaries]}")
print(f"  列数: {len(boundaries)-1}")
print("  [PASS]")

print("\n" + "=" * 70)
print("测试 4: 空间分析 - 构建表格")
print("=" * 70)
table = _build_table_from_grid(rows, boundaries, 1)
assert table is not None, "表格不应为 None"
print(f"  表格: {table.num_rows}×{table.num_cols}")
print(f"  单元格内容:")
for ri, row in enumerate(table.cells):
    texts = [c.text for c in row]
    print(f"    行 {ri+1}: {texts}")
assert table.num_rows == 3
assert table.num_cols == 3
print("  [PASS]")

print("\n" + "=" * 70)
print("测试 5: detect_tables_spatial 完整流程")
print("=" * 70)
tables = detect_tables_spatial(lines, PAGE_W, PAGE_H, 1)
print(f"  检测到 {len(tables)} 个表格")
assert len(tables) == 1
t = tables[0]
print(f"  表格尺寸: {t.num_rows}×{t.num_cols}")
print(f"  表格 bbox: {t.bbox}")
for ri, row in enumerate(t.cells):
    texts = [c.text for c in row]
    print(f"    行 {ri+1}: {texts}")
print("  [PASS]")

print("\n" + "=" * 70)
print("测试 6: 非表格数据 - 不应误检")
print("=" * 70)
# 模拟普通段落文本（连续的不对齐行）
non_table_lines = [
    OCRLine(text="第一章", score=0.9, bbox=(100, 100, 400, 140), polygon=[]),
    OCRLine(text="这是一个很长的段落文本，用于测试不会误判为表格。", score=0.9, bbox=(100, 200, 1800, 240), polygon=[]),
    OCRLine(text="这是第二行正文内容，连续的文字排列。", score=0.9, bbox=(100, 300, 1700, 340), polygon=[]),
    OCRLine(text="第三行继续正文", score=0.9, bbox=(100, 400, 1650, 440), polygon=[]),
]
tables = detect_tables_spatial(non_table_lines, PAGE_W, PAGE_H, 1)
print(f"  非表格文本检测结果: {len(tables)} 个表格 (预期: 0)")
assert len(tables) == 0, f"不应误检为表格, 实际检测到 {len(tables)}"
print("  [PASS]")

print("\n" + "=" * 70)
print("测试 7: Word 表格写入功能")
print("=" * 70)
from src.word_writer import _write_table_to_word
from docx import Document

doc = Document()
table_data = TableData(
    cells=[
        [CellData(text="序号"), CellData(text="名称"), CellData(text="备注")],
        [CellData(text="1"), CellData(text="设备A"), CellData(text="正常")],
        [CellData(text="2"), CellData(text="设备B"), CellData(text="异常")],
    ],
    bbox=(100, 100, 1500, 340),
    num_rows=3,
    num_cols=3,
    page_num=1,
)
_write_table_to_word(doc, table_data, {"size": 10, "latin": "Times New Roman", "east_asian": "宋体"}, PAGE_W, PAGE_H)
print(f"  Word 表格写入成功")
print(f"  Word 文档段落数: {len(doc.paragraphs)}")
print(f"  Word 表格数: {len(doc.tables)}")
assert len(doc.tables) == 1, f"应有 1 个表格, 实际 {len(doc.tables)}"
t = doc.tables[0]
print(f"  表格内容:")
for ri in range(len(t.rows)):
    texts = [t.cell(ri, ci).text for ci in range(len(t.columns))]
    print(f"    行 {ri+1}: {texts}")
print("  [PASS]")

print("\n" + "=" * 70)
print("测试 8: PageResult 含表格数据")
print("=" * 70)
# 生成新表格数据，不依赖之前作用域
table_result = detect_tables_spatial(lines, PAGE_W, PAGE_H, 1)
pr = PageResult(page_num=1, lines=lines, tables=table_result)
print(f"  PageResult.tables: {len(pr.tables)} 个表格")
assert len(pr.tables) == 1, f"预期 1, 实际 {len(pr.tables)}"
t = pr.tables[0]
print(f"  表格尺寸: {t.num_rows}×{t.num_cols}")
for ri, row in enumerate(t.cells):
    texts = [c.text for c in row]
    print(f"    行 {ri+1}: {texts}")
print("  [PASS]")

print("\n" + "=" * 70)
print("全部表格检测测试通过！")
print("=" * 70)